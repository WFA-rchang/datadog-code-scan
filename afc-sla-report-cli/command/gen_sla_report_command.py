import click
from docx import Document
from docx.shared import Pt
from gql import Client, gql
from datetime import datetime
from datetime import timezone
from docx.enum.text import WD_ALIGN_PARAGRAPH
from gql.transport.requests import RequestsHTTPTransport


@click.command(name='gen-sla-report', help="Generate AFC Quarterly SLA Availability Report")
@click.option("--date-start", type=click.DateTime(formats=["%Y-%m-%d"]), required=True, help="Start Date of the SLA Report [2024-01-01]")
@click.option("--date-end", type=click.DateTime(formats=["%Y-%m-%d"]), required=True, help="End Date of the SLA Report [2024-01-01]")
@click.option("--metrics-gateway-url", default="http://localhost:8000/graphql", show_default=True)
def gen_sla_report_command(date_start, date_end, metrics_gateway_url):
    # Set the start date to UTC
    date_start = date_start.replace(tzinfo=timezone.utc)
    # Set the end date to UTC
    date_end = date_end.replace(tzinfo=timezone.utc)

    # Input validation
    if date_start > date_end:
        raise click.ClickException("Start date must be before end date.")
    if ((date_end.year - date_start.year) * 12 + date_end.month - date_start.month) > 3:
        raise click.ClickException("Date range must be less than 3 months")
    if date_end > datetime.now(tz=timezone.utc):
        raise click.ClickException("To date must be earlier than current date")
    if date_start > datetime.now(tz=timezone.utc):
        raise click.ClickException("From date must be earlier than current date")

    # Load SLA Report template
    doc = Document(docx='sla_report_template.docx')

    # Get the Title table
    title_table = doc.tables[0]

    # Update the title date
    click.echo("Update Title Date...")
    title_table.cell(0, 1).text = f"From: {date_start.strftime('%m/%d/%Y')}\nTo: {date_end.strftime('%m/%d/%Y')}"
    title_table.cell(0, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    # title_table.cell(0, 1).paragraphs[0].style.font.name = 'Open Sans'
    title_table.cell(0, 1).paragraphs[0].style.font.size = Pt(11)
    click.echo("Update Title Date is done!")

    # calling Metrics-Gateway API to get the data
    click.echo("Calling Metrics-Gateway API...")
    transport = RequestsHTTPTransport(
        url=metrics_gateway_url,
        verify=False,
        retries=3,
    )
    client = Client(transport=transport, fetch_schema_from_transport=True)
    query = gql(
        f"""
        {{
            quarterlyServiceAvailabilityReport(fromDate: "{date_start.isoformat()}", toDate: "{date_end.isoformat()}") {{
                availabilities {{
                    date
                    availability
                    scheduledMaintenanceTotal
                    resultingDowntimeOfScheduledMaintenanceTotal
                    emergencyMaintenanceTotal
                    resultingDowntimeOfEmergencyMaintenanceTotal
                    totalDowntime
                }}
                incidents {{
                    fromDate
                    toDate
                    rootCause
                    salesforceCaseId
                }}
            }}
        }}
        """
    )

    sla_quarterly_report_availabilities = []
    sla_quarterly_report_incidents = []
    try:
        result = client.execute(query)
        sla_quarterly_report_response = result.get("quarterlyServiceAvailabilityReport")
        sla_quarterly_report_availabilities = sla_quarterly_report_response.get("availabilities")
        sla_quarterly_report_incidents = sla_quarterly_report_response.get("incidents")
    except Exception as e:
        click.echo(e)
        return

    # Get the Service Availability and Down Time table
    service_availability_table = doc.tables[1]

    # Update Service Availability and Down Time Table
    if len(sla_quarterly_report_availabilities) != 3:
        raise click.ClickException("SLA Report must have 3 availabilities data")

    click.echo("Update Service Availability and Down Time...")
    for i in range(3):
        sla_quarterly_report_availability_object = sla_quarterly_report_availabilities[i]
        sla_quarterly_report_availability_date = sla_quarterly_report_availability_object.get("date")
        sla_quarterly_report_availability_date_datetime = datetime.fromisoformat(sla_quarterly_report_availability_date)
        sla_quarterly_report_availability_date_month_string = sla_quarterly_report_availability_date_datetime.strftime("%B")
        sla_quarterly_report_availability = sla_quarterly_report_availability_object.get("availability")
        sla_quarterly_report_scheduled_maintenance_total = sla_quarterly_report_availability_object.get("scheduledMaintenanceTotal")
        sla_quarterly_report_resulting_downtime_of_scheduled_maintenance_total = sla_quarterly_report_availability_object.get("resultingDowntimeOfScheduledMaintenanceTotal")
        sla_quarterly_report_emergency_maintenance_total = sla_quarterly_report_availability_object.get("emergencyMaintenanceTotal")
        sla_quarterly_report_resulting_downtime_of_emergency_maintenance_total = sla_quarterly_report_availability_object.get("resultingDowntimeOfEmergencyMaintenanceTotal")
        sla_quarterly_report_total_downtime = sla_quarterly_report_availability_object.get("totalDowntime")

        service_availability_table.cell(i+1, 0).paragraphs[0].text = sla_quarterly_report_availability_date_month_string
        service_availability_table.cell(i+1, 1).paragraphs[0].text = f"{sla_quarterly_report_availability}%"
        service_availability_table.cell(i+1, 2).paragraphs[0].text = f"{sla_quarterly_report_scheduled_maintenance_total} / {sla_quarterly_report_resulting_downtime_of_scheduled_maintenance_total}"
        service_availability_table.cell(i+1, 3).paragraphs[0].text = f"{sla_quarterly_report_emergency_maintenance_total} / {sla_quarterly_report_resulting_downtime_of_emergency_maintenance_total}"
        service_availability_table.cell(i+1, 4).paragraphs[0].text = f"{sla_quarterly_report_total_downtime}"
    click.echo("Update Service Availability and Down Time is done!")

    # Get Incident description paragraph
    incident_description_paragraph = doc.paragraphs[7]

    # Get the Incident table
    incident_table = doc.tables[2]

    # There are incidents
    if len(sla_quarterly_report_incidents) > 0:
        # Clear the Incident Description paragraph
        click.echo("Clear Incident Description...")
        incident_description_paragraph_element = incident_description_paragraph._element
        incident_description_paragraph_element.getparent().remove(incident_description_paragraph_element)
        incident_description_paragraph_element._p = incident_description_paragraph_element._element = None
        click.echo("Clear Incident Description is done!")

        # Update the Incident table
        click.echo("Update Incident table...")
        for i in range(len(sla_quarterly_report_incidents)):
            sla_quarterly_report_incident_object = sla_quarterly_report_incidents[i]
            sla_quarterly_report_incident_from_date = sla_quarterly_report_incident_object.get("fromDate")
            sla_quarterly_report_incident_from_date_datetime = datetime.fromisoformat(sla_quarterly_report_incident_from_date)
            sla_quarterly_report_incident_to_date = sla_quarterly_report_incident_object.get("toDate")
            sla_quarterly_report_incident_to_date_datetime = datetime.fromisoformat(sla_quarterly_report_incident_to_date)
            sla_quarterly_report_incident_root_cause = sla_quarterly_report_incident_object.get("rootCause")
            sla_quarterly_report_incident_salesforce_case_id = sla_quarterly_report_incident_object.get("salesforceCaseId")
            if i < 3:
                incident_table.cell(i+1, 0).paragraphs[0].text = f"{sla_quarterly_report_incident_salesforce_case_id}"
                incident_table.cell(i+1, 1).paragraphs[0].text = f"{sla_quarterly_report_incident_from_date_datetime.strftime('%m/%d/%Y')}"
                incident_table.cell(i+1, 2).paragraphs[0].text = f"{sla_quarterly_report_incident_from_date_datetime.strftime('%H:%M')}"
                incident_table.cell(i+1, 3).paragraphs[0].text = f"{sla_quarterly_report_incident_to_date_datetime.strftime('%H:%M')}"
                incident_table.cell(i+1, 4).paragraphs[0].text = f"{sla_quarterly_report_incident_root_cause}"
            else:
                incident_table_new_row = incident_table.add_row()
                incident_table_new_row.cells[0].paragraphs[0].text = f"{sla_quarterly_report_incident_salesforce_case_id}"
                incident_table_new_row.cells[1].paragraphs[0].text = f"{sla_quarterly_report_incident_from_date_datetime.strftime('%m/%d/%Y')}"
                incident_table_new_row.cells[2].paragraphs[0].text = f"{sla_quarterly_report_incident_from_date_datetime.strftime('%H:%M')}"
                incident_table_new_row.cells[3].paragraphs[0].text = f"{sla_quarterly_report_incident_to_date_datetime.strftime('%H:%M')}"
                incident_table_new_row.cells[4].paragraphs[0].text = f"{sla_quarterly_report_incident_root_cause}"
        click.echo("Update Incident table is done!")
    else:
        click.echo("No Incidents found!")
        # Remove the Incident table
        click.echo("Remove Incident table...")
        incident_table._element.getparent().remove(incident_table._element)
        click.echo("Remove Incident table is done!")

    # Save the document
    doc.save('sla_report_generated.docx')
    click.echo("sla_report_generated.docx generated!")
